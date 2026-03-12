from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _h(doc: Document, text: Any, level: int = 1) -> None:
    if text:
        doc.add_heading(str(text).strip(), level=level)


def _p(doc: Document, text: Any) -> None:
    if text is None:
        return
    text = str(text).strip()
    if text:
        doc.add_paragraph(text)


def _bullets(doc: Document, items: Any) -> None:
    if not isinstance(items, list):
        return
    for item in items:
        text = item.get("text") or item.get("value") or item.get("item") or "" \
               if isinstance(item, dict) else str(item)
        text = str(text).strip()
        if text:
            doc.add_paragraph(text, style="List Bullet")


def _kv_table(doc: Document, rows: list[dict], col1: str = "Metric", col2: str = "Value") -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = col1
    hdr[1].text = col2
    for row in rows:
        if not isinstance(row, dict):
            continue
        cells = table.add_row().cells
        cells[0].text = str(row.get("metric", row.get("country", "")))
        cells[1].text = str(row.get("details", row.get("value", row.get("cagr", ""))))


def render_docx(article: dict, output_path: Path) -> None:
    doc = Document()
    market = article.get("market_name", "")
    s = article.get("sections", {})

    # ── Title block
    tb = s.get("title_block", {})
    _h(doc, tb.get("h1_title", market), level=0)
    seg_line = tb.get("segmentation_line", "")
    period = tb.get("forecast_period", "")
    if seg_line:
        _p(doc, f"{seg_line}. Forecast for {period}." if period else seg_line)

    # ── Section 1: Opening block
    ob = s.get("opening_block", {})
    _h(doc, ob.get("heading", ""), level=1)
    _p(doc, ob.get("paragraph", ""))

    # ── Section 2: Summary
    sm = s.get("summary", {})
    _h(doc, sm.get("heading", ""), level=1)

    _h(doc, sm.get("key_drivers_heading", "Key Drivers"), level=2)
    _bullets(doc, sm.get("key_drivers", []))

    _h(doc, sm.get("key_segments_heading", "Key Segments Analyzed in the Report"), level=2)
    _bullets(doc, sm.get("key_segments", []))

    _h(doc, sm.get("analyst_opinion_heading", "Analyst Opinion"), level=2)
    _p(doc, sm.get("analyst_opinion", ""))

    # ── Section 3: Market Definition
    md = s.get("market_definition", {})
    _h(doc, md.get("heading", ""), level=1)
    _p(doc, md.get("paragraph", ""))

    # ── Section 4: Inclusions
    inc = s.get("inclusions", {})
    _h(doc, inc.get("heading", ""), level=1)
    _bullets(doc, inc.get("bullets", []))

    # ── Section 5: Exclusions
    exc = s.get("exclusions", {})
    _h(doc, exc.get("heading", ""), level=1)
    _bullets(doc, exc.get("bullets", []))

    # ── Section 6: Methodology
    meth = s.get("methodology", {})
    _h(doc, meth.get("heading", ""), level=1)
    _p(doc, meth.get("paragraph", ""))

    # ── Section 7: Drivers, Restraints, Trends
    dyn = s.get("market_dynamics", {})
    _h(doc, dyn.get("heading", ""), level=1)

    _h(doc, dyn.get("drivers_heading", "Drivers"), level=2)
    val = dyn.get("drivers_paragraph", "")
    if isinstance(val, list):
        _bullets(doc, val)
    else:
        _p(doc, val)

    _h(doc, dyn.get("restraints_heading", "Restraints"), level=2)
    val = dyn.get("restraints_paragraph", "")
    if isinstance(val, list):
        _bullets(doc, val)
    else:
        _p(doc, val)

    _h(doc, dyn.get("trends_heading", "Trends"), level=2)
    val = dyn.get("trends_paragraph", "")
    if isinstance(val, list):
        _bullets(doc, val)
    else:
        _p(doc, val)

    # ── Section 8: Segmental Analysis
    sa = s.get("segmental_analysis", {})
    _h(doc, sa.get("heading", "Segmental Analysis"), level=1)
    for item in sa.get("items", []):
        if not isinstance(item, dict):
            continue
        _h(doc, item.get("subheading", ""), level=2)
        val = item.get("paragraph", "")
        if isinstance(val, list):
            _bullets(doc, val)
        else:
            _p(doc, val)

    # ── Section 9: Competitive Aligners
    ca = s.get("competitive_aligners", {})
    _h(doc, ca.get("heading", "Competitive Aligners for Market Players"), level=1)
    _p(doc, ca.get("paragraph_1", ""))
    _p(doc, ca.get("paragraph_2", ""))

    # ── Section 10: Key Players
    kp = s.get("key_players", {})
    _h(doc, kp.get("heading", "Key Players"), level=1)
    _bullets(doc, kp.get("items", []))

    # ── Section 11: Strategic Outlook
    so = s.get("strategic_outlook", {})
    _h(doc, so.get("heading", ""), level=1)
    _p(doc, so.get("paragraph", ""))

    # ── Section 12: Scope of the Report
    sr = s.get("scope_of_report", {})
    _h(doc, sr.get("heading", "Scope of the Report"), level=1)
    _kv_table(doc, sr.get("table_rows", []))

    # ── Section 13: Bibliography
    bib = s.get("bibliography", {})
    _h(doc, bib.get("heading", "Bibliography"), level=1)
    _bullets(doc, bib.get("items", []))

    # ── Section 14: FAQs
    faqs = s.get("faqs", {})
    _h(doc, faqs.get("heading", "Frequently Asked Questions"), level=1)
    for item in faqs.get("items", []):
        if isinstance(item, dict):
            q = item.get("question", "")
            a = item.get("answer", "")
            if q:
                p = doc.add_paragraph()
                run = p.add_run(q)
                run.bold = True
            if a:
                doc.add_paragraph(a)
        else:
            _p(doc, str(item))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
